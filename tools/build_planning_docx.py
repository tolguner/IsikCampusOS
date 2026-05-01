from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "bitirme-projesi-planlama-dokumani.md"
OUTPUT = ROOT / "docs" / "IsikCampusOS_Bitirme_Projesi_Planlama_Dokumani.docx"


ACCENT = "1F4E79"
LIGHT = "EAF2F8"
TEXT = RGBColor(31, 41, 55)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text.strip())
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D7DEE8", size="6"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    for name, size in (("Heading 1", 17), ("Heading 2", 13), ("Heading 3", 11.5)):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(ACCENT)
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(5)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "IsikCampusOS | Bitirme Projesi Planlama Dokümanı"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if header.runs:
        header.runs[0].font.size = Pt(8)
        header.runs[0].font.color.rgb = RGBColor(107, 114, 128)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Danışman değerlendirmesi için hazırlanmıştır")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(107, 114, 128)


def add_cover(doc):
    for _ in range(2):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("IsikCampusOS")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor.from_string(ACCENT)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Bitirme Projesi Planlama Dokümanı")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(55, 65, 81)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(16)
    run = note.add_run(
        "Mikroservis mimarili dijital kampüs yönetim platformu için proje kapsamı, "
        "teknoloji yığını, mimari yaklaşım, yol haritası ve başarı kriterleri"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(75, 85, 99)

    doc.add_paragraph()
    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = True
    rows = [
        ("Proje türü", "Bitirme projesi / yazılım mühendisliği uygulaması"),
        ("Mimari yaklaşım", "Microservices, API Gateway, event-driven communication"),
        ("Ana teknoloji yığını", "React, TypeScript, Java 21, Spring Boot, PostgreSQL, Kafka, Docker"),
        ("Hazırlanma tarihi", "29 Nisan 2026"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        set_cell_text(row.cells[0], label, bold=True, color=ACCENT)
        set_cell_text(row.cells[1], value)
        set_cell_shading(row.cells[0], LIGHT)
    set_table_borders(meta)

    doc.add_page_break()


def add_section_index(doc, headings):
    doc.add_heading("İçindekiler", level=1)
    for text in headings:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.add_run(text)
    doc.add_page_break()


def add_markdown_table(doc, lines):
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)

    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_table_borders(table)

    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            set_cell_text(cell, value, bold=(r_idx == 0), color=("FFFFFF" if r_idx == 0 else None))
            if r_idx == 0:
                set_cell_shading(cell, ACCENT)
            elif r_idx % 2 == 1:
                set_cell_shading(cell, "F8FAFC")

    doc.add_paragraph()


def add_inline_text(paragraph, text):
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def build():
    markdown = SOURCE.read_text(encoding="utf-8").splitlines()
    headings = [line[3:].strip() for line in markdown if line.startswith("## ") and not line.startswith("### ")]

    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_cover(doc)
    add_section_index(doc, headings)

    table_buffer = []
    skip_title = True

    def flush_table():
        nonlocal table_buffer
        if table_buffer:
            add_markdown_table(doc, table_buffer)
            table_buffer = []

    for line in markdown:
        if skip_title and line.startswith("# "):
            skip_title = False
            continue

        if line.strip().startswith("|"):
            table_buffer.append(line)
            continue

        flush_table()
        stripped = line.strip()

        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            add_inline_text(p, stripped[2:])
        else:
            p = doc.add_paragraph()
            add_inline_text(p, stripped)

    flush_table()

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Calibri"

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
